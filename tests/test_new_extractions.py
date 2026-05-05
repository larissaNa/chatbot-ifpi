import os
import sys
from io import BytesIO

# Adiciona o diretório raiz ao path para poder importar os módulos do app
sys.path.append(os.getcwd())

from apps.core.documents.agents.extraction_agent import _extract_from_docx, _extract_from_txt

def test_txt_extraction():
    print("Testando extração de TXT...")
    content = "Este é um teste de extração de texto de um arquivo TXT.".encode("utf-8")
    texto, paginas, tipo, obs = _extract_from_txt(content)
    print(f"Texto: {texto}")
    print(f"Páginas: {paginas}")
    print(f"Tipo: {tipo}")
    print(f"Obs: {obs}")
    assert "Este é um teste" in texto
    assert tipo == "TXT_TEXTUAL"
    print("Sucesso!\n")

def test_docx_extraction():
    print("Testando extração de DOCX...")
    import docx
    doc = docx.Document()
    doc.add_heading('Título do Documento', 0)
    doc.add_paragraph('Este é um parágrafo de teste em DOCX.')
    
    f = BytesIO()
    doc.save(f)
    content = f.getvalue()
    
    texto, paginas, tipo, obs = _extract_from_docx(content)
    print(f"Texto: {texto}")
    print(f"Páginas: {paginas}")
    print(f"Tipo: {tipo}")
    print(f"Obs: {obs}")
    assert "Este é um parágrafo de teste" in texto
    assert tipo == "DOCX_TEXTUAL"
    print("Sucesso!\n")

if __name__ == "__main__":
    try:
        test_txt_extraction()
        test_docx_extraction()
        print("Todos os testes de extração passaram!")
    except Exception as e:
        print(f"Erro nos testes: {e}")
        sys.exit(1)
