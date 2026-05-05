import re
import uuid
import logging
import os

import fitz
import docx
from bs4 import BeautifulSoup
from langchain_core.tools import tool

from .download_utils import baixar_arquivo_resiliente
from apps.core.utils.fetcher import fetch_url_content

logger = logging.getLogger(__name__)

def _normalize_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [l.strip() for l in text.split("\n")]
    cleaned = []
    for line in lines:
        if not line:
            if cleaned and cleaned[-1] != "":
                cleaned.append("")
        else:
            line = re.sub(r"\s+", " ", line)
            cleaned.append(line)
    paragraphs = [l for l in cleaned if l]
    return "\n\n".join(paragraphs)


def _detect_language(text: str) -> str:
    if not text:
        return "desconhecido"
    lower = text.lower()
    markers = [
        " que ",
        " de ",
        " para ",
        " nao ",
        "cao",
        "coes",
        " nº ",
        " art. ",
    ]
    score = 0
    for marker in markers:
        if marker in lower:
            score += 1
    if score >= 2:
        return "pt"
    return "desconhecido"


def _format_table_as_markdown(table) -> str:
    """Converte um objeto Table do PyMuPDF para Markdown e Texto Corrido (Verbalizado)."""
    data = table.extract()
    if not data:
        return ""
    
    # 1. Gerar Markdown
    md = "\n"
    headers = []
    for i, row in enumerate(data):
        row_cleaned = [str(cell).replace("\n", " ").strip() if cell is not None else "" for cell in row]
        if not any(row_cleaned):
            continue
        if i == 0:
            headers = row_cleaned
        md += "| " + " | ".join(row_cleaned) + " |\n"
        if i == 0:
            md += "| " + " | ".join(["---"] * len(row_cleaned)) + " |\n"
    
    # 2. Gerar Texto Corrido (Verbalização para melhor embedding)
    verbalized = ""
    if len(data) > 1 and headers:
        verbalized = "\nDescrição da tabela:\n"
        verbalized += f"Esta tabela contém informações sobre: {', '.join([h for h in headers if h])}.\n"
        for i, row in enumerate(data[1:], start=1):
            row_parts = []
            for j, cell in enumerate(row):
                if j < len(headers) and headers[j]:
                    val = str(cell).strip() if cell else "não informado"
                    row_parts.append(f"{headers[j]} é '{val}'")
            if row_parts:
                verbalized += f"- Registro {i}: {', '.join(row_parts)}.\n"
    
    return md + verbalized + "\n"


def _extract_from_pdf_url(url: str):
    observacoes = []
    try:
        content = baixar_arquivo_resiliente(url)
    except Exception as e:
        return "", 0, "PDF_TEXTUAL", [f"Erro ao baixar PDF resiliente: {e}"]
        
    try:
        pdf = fitz.open(stream=content, filetype="pdf")
    except Exception as e:
        return "", 0, "PDF_TEXTUAL", [f"Erro ao abrir PDF: {e}"]
    
    texts = []
    lengths = []
    
    for page in pdf:
        # Tenta encontrar tabelas na página
        try:
            tabs = page.find_tables()
            tables = tabs.tables
        except Exception:
            tables = []
            
        # 1. Identificar bboxes das tabelas
        table_bboxes = [t.bbox for t in tables]
        
        # 2. Obter blocos de texto
        blocks = page.get_text("blocks")
        # bloco: (x0, y0, x1, y1, "texto", block_no, block_type)
        
        page_elements = []
        
        # Adiciona blocos de texto que NÃO estão dentro de tabelas
        for b in blocks:
            if b[6] == 0: # tipo texto
                block_rect = fitz.Rect(b[:4])
                is_inside_table = False
                for t_bbox in table_bboxes:
                    # Se o bloco está majoritariamente dentro da tabela, ignoramos
                    if block_rect.intersect(t_bbox).get_area() > (block_rect.get_area() * 0.5):
                        is_inside_table = True
                        break
                if not is_inside_table:
                    # Normalização leve do bloco de texto
                    content = b[4].strip()
                    if content:
                        # Substitui quebras de linha internas por espaços para manter o parágrafo
                        content = re.sub(r"\s+", " ", content)
                        page_elements.append({"type": "text", "y": b[1], "content": content})
        
        # Adiciona as tabelas formatadas (já estão em MD + verbalizado)
        for t in tables:
            md_table = _format_table_as_markdown(t)
            if md_table.strip():
                page_elements.append({"type": "table", "y": t.bbox[1], "content": md_table})
        
        # Ordena tudo pelo eixo Y (vertical) para manter a ordem de leitura
        page_elements.sort(key=lambda x: x["y"])
        
        page_text = "\n\n".join([el["content"] for el in page_elements])
        texts.append(page_text)
        lengths.append(len(page_text.strip()))

    raw_text = "\n\n".join(texts)
    # Atribuímos diretamente ao texto final, pois a limpeza já foi feita por bloco
    texto = raw_text
    
    num_pages = len(pdf)
    pages_with_text = sum(1 for length in lengths if length >= 20)
    if num_pages == 0:
        tipo_extracao = "PDF_TEXTUAL"
        observacoes.append("PDF vazio ou sem paginas.")
    else:
        ratio = pages_with_text / float(num_pages)
        if ratio >= 0.5:
            tipo_extracao = "PDF_TEXTUAL"
        else:
            tipo_extracao = "PDF_OCR"
            observacoes.append("PDF provavelmente escaneado. OCR nao aplicado ou incompleto.")
    
    # Validação obrigatória de conteúdo extraído
    if not texto or len(texto.strip()) < 100:
        logger.error(f"[EXTRACTOR] Falha na extração de texto do PDF {url}: Conteúdo insuficiente ({len(texto) if texto else 0} caracteres)")
        return "", 0, "PDF_TEXTUAL", ["Falha na extração de texto do PDF ou conteúdo insuficiente (< 100 caracteres)."]

    return texto, num_pages, tipo_extracao, observacoes


def _extract_from_docx(content: bytes):
    """Extrai texto de um arquivo .docx usando python-docx."""
    from io import BytesIO
    try:
        doc = docx.Document(BytesIO(content))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # Também extrai tabelas
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        texto = "\n\n".join(full_text)
        # Normalização
        texto = _normalize_text(texto)
        paginas = max(1, int(len(texto) / 1800))
        return texto, paginas, "DOCX_TEXTUAL", []
    except Exception as e:
        return "", 0, "DOCX_TEXTUAL", [f"Erro ao extrair DOCX: {e}"]


def _extract_from_txt(content: bytes):
    """Extrai texto de um arquivo .txt."""
    try:
        # Tenta decodificar como utf-8, fallback para latin-1
        try:
            texto = content.decode("utf-8")
        except UnicodeDecodeError:
            texto = content.decode("latin-1")
            
        texto = _normalize_text(texto)
        paginas = max(1, int(len(texto) / 1800))
        return texto, paginas, "TXT_TEXTUAL", []
    except Exception as e:
        return "", 0, "TXT_TEXTUAL", [f"Erro ao extrair TXT: {e}"]


def _extract_from_html_url(url: str):
    observacoes = []
    # Usamos o novo fetcher robusto que suporta fallback e proxy
    try:
        html = fetch_url_content(url)
    except Exception as e:
        return None, "", 0, [f"Erro ao baixar HTML com fetcher robusto: {e}"]
        
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "nav", "header", "footer", "aside", "noscript"]):
        tag.decompose()
    titulo = None
    title_tag = soup.find("title")
    if title_tag:
        title_text = title_tag.get_text(strip=True)
        if title_text:
            titulo = title_text
    parts = []
    for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li"]):
        text = tag.get_text(" ", strip=True)
        if text:
            parts.append(text)
    if not parts:
        body_text = soup.get_text(" ", strip=True)
        if body_text:
            parts.append(body_text)
    raw_text = "\n\n".join(parts)
    texto = _normalize_text(raw_text)
    if texto:
        paginas_estimadas = max(1, int(len(texto) / 1800))
    else:
        paginas_estimadas = 0
        observacoes.append("Nenhum conteudo textual significativo encontrado no HTML.")
    return titulo, texto, paginas_estimadas, observacoes


@tool
def extrair_conteudo(analise: dict):
    """
    Recebe a analise de um link e extrai o conteudo textual e metadados.
    Suporta PDF direto, HTML com links para PDF e HTML textual.
    Retorna um dicionario com os documentos extraidos e indicacao do proximo agente.
    """
    fonte = analise.get("url") or analise.get("fonte") or ""
    tipo_conteudo = analise.get("tipo_conteudo")
    pdfs = analise.get("pdfs_encontrados") or []
    documentos = []
    observacoes = []
    tipo_extracao_global = None
    # Normaliza o tipo_conteudo para maiúsculas e remove espaços
    tipo_conteudo_norm = str(tipo_conteudo).upper().strip() if tipo_conteudo else ""
    
    if tipo_conteudo_norm in ["PDF_DIRETO", "PDF"]:
        texto, paginas, tipo_extracao, obs = _extract_from_pdf_url(fonte)
        if obs:
            observacoes.extend(obs)
        idioma = _detect_language(texto)
        doc = {
            "id_documento": str(uuid.uuid4()),
            "titulo": analise.get("titulo"),
            "texto": texto,
            "metadata": {
                "paginas_estimadas": paginas,
                "idioma": idioma or "desconhecido",
                "tipo_extracao": tipo_extracao,
            },
        }
        documentos.append(doc)
        tipo_extracao_global = tipo_extracao
    elif tipo_conteudo_norm in ["DOCX", "WORD"]:
        try:
            content = baixar_arquivo_resiliente(fonte)
            texto, paginas, tipo_extracao, obs = _extract_from_docx(content)
            if obs:
                observacoes.extend(obs)
            idioma = _detect_language(texto)
            doc = {
                "id_documento": str(uuid.uuid4()),
                "titulo": analise.get("titulo"),
                "texto": texto,
                "metadata": {
                    "paginas_estimadas": paginas,
                    "idioma": idioma or "desconhecido",
                    "tipo_extracao": tipo_extracao,
                },
            }
            documentos.append(doc)
            tipo_extracao_global = tipo_extracao
        except Exception as e:
            observacoes.append(f"Erro ao processar DOCX: {e}")
            tipo_extracao_global = "DOCX_TEXTUAL"
    elif tipo_conteudo_norm == "TXT":
        try:
            content = baixar_arquivo_resiliente(fonte)
            texto, paginas, tipo_extracao, obs = _extract_from_txt(content)
            if obs:
                observacoes.extend(obs)
            idioma = _detect_language(texto)
            doc = {
                "id_documento": str(uuid.uuid4()),
                "titulo": analise.get("titulo"),
                "texto": texto,
                "metadata": {
                    "paginas_estimadas": paginas,
                    "idioma": idioma or "desconhecido",
                    "tipo_extracao": tipo_extracao,
                },
            }
            documentos.append(doc)
            tipo_extracao_global = tipo_extracao
        except Exception as e:
            observacoes.append(f"Erro ao processar TXT: {e}")
            tipo_extracao_global = "TXT_TEXTUAL"
    elif tipo_conteudo_norm == "HTML_COM_PDF":
        if not pdfs:
            observacoes.append("tipo_conteudo=HTML_COM_PDF mas pdfs_encontrados esta vazio.")
        for item in pdfs:
            pdf_url = item.get("url")
            if not pdf_url:
                continue
            texto, paginas, tipo_extracao, obs = _extract_from_pdf_url(pdf_url)
            if obs:
                observacoes.append(f"{pdf_url}: " + " ".join(obs))
            idioma = _detect_language(texto)
            doc = {
                "id_documento": str(uuid.uuid4()),
                "titulo": item.get("descricao") or analise.get("titulo"),
                "texto": texto,
                "metadata": {
                    "paginas_estimadas": paginas,
                    "idioma": idioma or "desconhecido",
                    "tipo_extracao": tipo_extracao,
                },
            }
            documentos.append(doc)
            if not tipo_extracao_global:
                tipo_extracao_global = tipo_extracao
            elif tipo_extracao_global == "PDF_TEXTUAL" and tipo_extracao == "PDF_OCR":
                tipo_extracao_global = "PDF_OCR"
        if not tipo_extracao_global:
            tipo_extracao_global = "PDF_TEXTUAL"
    elif tipo_conteudo_norm in ["HTML_TEXTO", "PAGINA"]:
        titulo, texto, paginas, obs = _extract_from_html_url(fonte)
        if obs:
            observacoes.extend(obs)
        idioma = _detect_language(texto)
        doc = {
            "id_documento": str(uuid.uuid4()),
            "titulo": titulo or analise.get("titulo"),
            "texto": texto,
            "metadata": {
                "paginas_estimadas": paginas,
                "idioma": idioma or "desconhecido",
                "tipo_extracao": "HTML",
            },
        }
        documentos.append(doc)
        tipo_extracao_global = "HTML"
    else:
        observacoes.append(f"tipo_conteudo invalido ou nao suportado: {tipo_conteudo_norm}")
        tipo_extracao_global = "HTML"
    return {
        "fonte": fonte,
        "tipo_extracao": tipo_extracao_global or "HTML",
        "documentos": documentos,
        "observacoes": " ".join(observacoes).strip(),
        "proximo_agente": "AGENTE_PROCESSAMENTO",
    }